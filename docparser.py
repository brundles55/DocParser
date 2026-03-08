#!/usr/bin/env python3
"""
DocParser - Document to Structured Data Converter
Converts PDFs and Office documents into clean, structured Markdown
optimized for LLM consumption and RAG pipelines.

Author: Built with Claude for Josh's infrastructure
License: MIT
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Optional

import yaml

# --- Document Models ---

@dataclass
class TableData:
    """Represents an extracted table."""
    headers: list[str]
    rows: list[list[str]]
    caption: str = ""

    def to_markdown(self) -> str:
        if not self.headers and not self.rows:
            return ""
        lines = []
        if self.caption:
            lines.append(f"*{self.caption}*\n")
        
        # Use headers if available, otherwise generate from first row
        if self.headers:
            hdrs = self.headers
        elif self.rows:
            hdrs = [f"Col {i+1}" for i in range(len(self.rows[0]))]
        else:
            return ""
        
        lines.append("| " + " | ".join(hdrs) + " |")
        lines.append("| " + " | ".join(["---"] * len(hdrs)) + " |")
        for row in self.rows:
            # Pad row if shorter than headers
            padded = row + [""] * (len(hdrs) - len(row))
            lines.append("| " + " | ".join(padded[:len(hdrs)]) + " |")
        return "\n".join(lines)


@dataclass
class Section:
    """Represents a document section with heading hierarchy."""
    level: int  # 1-6 for h1-h6
    title: str
    content: str
    subsections: list = field(default_factory=list)
    tables: list = field(default_factory=list)


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""
    filename: str
    filepath: str
    file_type: str
    file_size_bytes: int
    page_count: int = 0
    title: str = ""
    author: str = ""
    created_date: str = ""
    modified_date: str = ""
    subject: str = ""
    keywords: list = field(default_factory=list)
    word_count: int = 0
    char_count: int = 0
    sha256: str = ""
    parsed_at: str = ""


@dataclass
class ParsedDocument:
    """Complete parsed document ready for export."""
    metadata: DocumentMetadata
    sections: list  # list of Section
    raw_text: str = ""
    tables: list = field(default_factory=list)  # standalone tables
    warnings: list = field(default_factory=list)


# --- PDF Parser ---

class PDFParser:
    """Extracts structured content from PDF files using PyMuPDF."""

    def parse(self, filepath: str) -> ParsedDocument:
        import fitz  # PyMuPDF

        doc = fitz.open(filepath)
        path = Path(filepath)
        
        # Extract metadata
        meta = doc.metadata or {}
        file_stat = path.stat()
        
        # Compute hash
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
        
        all_text_blocks = []
        all_tables = []
        warnings = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text blocks with position info
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            
            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block.get("lines", []):
                        text = ""
                        max_font_size = 0
                        is_bold = False
                        for span in line.get("spans", []):
                            text += span["text"]
                            max_font_size = max(max_font_size, span["size"])
                            if "bold" in span.get("font", "").lower() or \
                               "Bold" in span.get("font", ""):
                                is_bold = True
                        
                        text = text.strip()
                        if text:
                            all_text_blocks.append({
                                "text": text,
                                "font_size": max_font_size,
                                "is_bold": is_bold,
                                "page": page_num + 1,
                                "y_pos": line["bbox"][1],
                            })
            
            # Attempt table extraction
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table in tables.tables:
                        extracted = table.extract()
                        if extracted and len(extracted) > 1:
                            headers = [str(c).strip() if c else "" for c in extracted[0]]
                            rows = []
                            for row in extracted[1:]:
                                rows.append([str(c).strip() if c else "" for c in row])
                            all_tables.append(TableData(
                                headers=headers,
                                rows=rows,
                                caption=f"Table from page {page_num + 1}"
                            ))
            except Exception as e:
                warnings.append(f"Table extraction failed on page {page_num + 1}: {e}")
        
        # Build sections from text blocks using font size heuristics
        sections = self._build_sections(all_text_blocks)
        raw_text = "\n".join(b["text"] for b in all_text_blocks)
        
        metadata = DocumentMetadata(
            filename=path.name,
            filepath=str(path.absolute()),
            file_type="pdf",
            file_size_bytes=file_stat.st_size,
            page_count=len(doc),
            title=meta.get("title", "") or path.stem,
            author=meta.get("author", ""),
            created_date=meta.get("creationDate", ""),
            modified_date=meta.get("modDate", ""),
            subject=meta.get("subject", ""),
            keywords=[k.strip() for k in meta.get("keywords", "").split(",") if k.strip()],
            word_count=len(raw_text.split()),
            char_count=len(raw_text),
            sha256=sha256,
            parsed_at=datetime.now().isoformat(),
        )
        
        doc.close()
        
        return ParsedDocument(
            metadata=metadata,
            sections=sections,
            raw_text=raw_text,
            tables=all_tables,
            warnings=warnings,
        )

    def _build_sections(self, blocks: list) -> list:
        """Use font size analysis to identify headings and build section hierarchy."""
        if not blocks:
            return []
        
        # Analyze font size distribution
        font_sizes = [b["font_size"] for b in blocks]
        if not font_sizes:
            return []
        
        avg_size = sum(font_sizes) / len(font_sizes)
        max_size = max(font_sizes)
        
        # Determine heading thresholds
        # Anything significantly larger than average is likely a heading
        h1_threshold = avg_size * 1.5
        h2_threshold = avg_size * 1.25
        h3_threshold = avg_size * 1.1
        
        sections = []
        current_section = None
        current_content_lines = []
        
        for block in blocks:
            text = block["text"]
            size = block["font_size"]
            is_bold = block["is_bold"]
            
            # Determine if this is a heading
            heading_level = 0
            if size >= h1_threshold or (size == max_size and is_bold):
                heading_level = 1
            elif size >= h2_threshold and is_bold:
                heading_level = 2
            elif (size >= h3_threshold and is_bold) or (is_bold and size > avg_size):
                heading_level = 3
            
            if heading_level > 0 and len(text) < 200:  # Headings shouldn't be super long
                # Save previous section
                if current_section or current_content_lines:
                    if current_section is None:
                        current_section = Section(level=1, title="", content="")
                    current_section.content = "\n".join(current_content_lines).strip()
                    sections.append(current_section)
                    current_content_lines = []
                
                current_section = Section(level=heading_level, title=text, content="")
            else:
                current_content_lines.append(text)
        
        # Don't forget the last section
        if current_content_lines:
            if current_section is None:
                current_section = Section(level=1, title="Document Content", content="")
            current_section.content = "\n".join(current_content_lines).strip()
            sections.append(current_section)
        elif current_section:
            sections.append(current_section)
        
        return sections


# --- DOCX Parser ---

class DOCXParser:
    """Extracts structured content from Word documents."""

    def parse(self, filepath: str) -> ParsedDocument:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(filepath)
        path = Path(filepath)
        file_stat = path.stat()
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
        
        # Extract core properties
        props = doc.core_properties
        
        sections = []
        tables = []
        raw_lines = []
        current_section = None
        current_content_lines = []
        warnings = []
        
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            
            if tag == "p":
                # Process paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                
                if para is None:
                    continue
                
                text = para.text.strip()
                if not text:
                    continue
                
                raw_lines.append(text)
                style_name = (para.style.name or "").lower() if para.style else ""
                
                # Detect heading level
                heading_level = 0
                if "heading" in style_name:
                    try:
                        heading_level = int(re.search(r'\d+', style_name).group())
                    except (AttributeError, ValueError):
                        heading_level = 1
                elif style_name == "title":
                    heading_level = 1
                elif style_name == "subtitle":
                    heading_level = 2
                
                if heading_level > 0:
                    # Save current section
                    if current_section or current_content_lines:
                        if current_section is None:
                            current_section = Section(level=1, title="", content="")
                        current_section.content = "\n".join(current_content_lines).strip()
                        sections.append(current_section)
                        current_content_lines = []
                    
                    current_section = Section(level=heading_level, title=text, content="")
                else:
                    # Check for list items
                    if style_name.startswith("list"):
                        text = f"- {text}"
                    current_content_lines.append(text)
            
            elif tag == "tbl":
                # Process table
                for table in doc.tables:
                    if table._element is element:
                        try:
                            headers = []
                            rows = []
                            for i, row in enumerate(table.rows):
                                cells = [cell.text.strip() for cell in row.cells]
                                if i == 0:
                                    headers = cells
                                else:
                                    rows.append(cells)
                            
                            td = TableData(headers=headers, rows=rows)
                            tables.append(td)
                            
                            # Also add table markdown to current content
                            current_content_lines.append("")
                            current_content_lines.append(td.to_markdown())
                            current_content_lines.append("")
                        except Exception as e:
                            warnings.append(f"Table extraction error: {e}")
                        break
        
        # Save last section
        if current_content_lines:
            if current_section is None:
                current_section = Section(level=1, title="Document Content", content="")
            current_section.content = "\n".join(current_content_lines).strip()
            sections.append(current_section)
        elif current_section:
            sections.append(current_section)
        
        raw_text = "\n".join(raw_lines)
        
        metadata = DocumentMetadata(
            filename=path.name,
            filepath=str(path.absolute()),
            file_type="docx",
            file_size_bytes=file_stat.st_size,
            page_count=0,  # DOCX doesn't have fixed pages
            title=props.title or path.stem,
            author=props.author or "",
            created_date=str(props.created) if props.created else "",
            modified_date=str(props.modified) if props.modified else "",
            subject=props.subject or "",
            keywords=[k.strip() for k in (props.keywords or "").split(",") if k.strip()],
            word_count=len(raw_text.split()),
            char_count=len(raw_text),
            sha256=sha256,
            parsed_at=datetime.now().isoformat(),
        )
        
        return ParsedDocument(
            metadata=metadata,
            sections=sections,
            raw_text=raw_text,
            tables=tables,
            warnings=warnings,
        )


# --- Plain Text / Markdown Parser ---

class TextParser:
    """Handles .txt and .md files with minimal transformation."""

    def parse(self, filepath: str) -> ParsedDocument:
        path = Path(filepath)
        file_stat = path.stat()
        content = path.read_text(encoding="utf-8", errors="replace")
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
        
        ext = path.suffix.lower()
        sections = []
        
        if ext == ".md":
            # Parse markdown headings
            current_section = None
            current_lines = []
            
            for line in content.split("\n"):
                heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
                if heading_match:
                    if current_section or current_lines:
                        if current_section is None:
                            current_section = Section(level=1, title="", content="")
                        current_section.content = "\n".join(current_lines).strip()
                        sections.append(current_section)
                        current_lines = []
                    
                    level = len(heading_match.group(1))
                    title = heading_match.group(2).strip()
                    current_section = Section(level=level, title=title, content="")
                else:
                    current_lines.append(line)
            
            if current_lines:
                if current_section is None:
                    current_section = Section(level=1, title="Document Content", content="")
                current_section.content = "\n".join(current_lines).strip()
                sections.append(current_section)
        else:
            sections.append(Section(
                level=1,
                title=path.stem,
                content=content.strip(),
            ))
        
        metadata = DocumentMetadata(
            filename=path.name,
            filepath=str(path.absolute()),
            file_type=ext.lstrip("."),
            file_size_bytes=file_stat.st_size,
            title=path.stem,
            word_count=len(content.split()),
            char_count=len(content),
            sha256=sha256,
            parsed_at=datetime.now().isoformat(),
        )
        
        return ParsedDocument(
            metadata=metadata,
            sections=sections,
            raw_text=content,
        )


# --- Output Formatters ---

class MarkdownFormatter:
    """Converts ParsedDocument to clean, LLM-optimized Markdown."""

    def format(self, doc: ParsedDocument, include_metadata: bool = True) -> str:
        lines = []
        
        if include_metadata:
            lines.append("---")
            lines.append(f"title: \"{doc.metadata.title}\"")
            lines.append(f"source_file: \"{doc.metadata.filename}\"")
            lines.append(f"file_type: {doc.metadata.file_type}")
            if doc.metadata.author:
                lines.append(f"author: \"{doc.metadata.author}\"")
            if doc.metadata.page_count:
                lines.append(f"pages: {doc.metadata.page_count}")
            lines.append(f"word_count: {doc.metadata.word_count}")
            lines.append(f"parsed_at: {doc.metadata.parsed_at}")
            lines.append(f"sha256: {doc.metadata.sha256[:16]}...")
            lines.append("---")
            lines.append("")
        
        # Write sections
        for section in doc.sections:
            if section.title:
                lines.append(f"{'#' * section.level} {section.title}")
                lines.append("")
            if section.content:
                # Clean up content - collapse excessive whitespace
                content = re.sub(r'\n{3,}', '\n\n', section.content)
                lines.append(content)
                lines.append("")
        
        # Append standalone tables not already in sections
        if doc.tables:
            has_tables_in_sections = any(
                "| " in s.content for s in doc.sections if s.content
            )
            if not has_tables_in_sections:
                lines.append("## Extracted Tables")
                lines.append("")
                for i, table in enumerate(doc.tables):
                    if table.caption:
                        lines.append(f"### {table.caption}")
                    lines.append(table.to_markdown())
                    lines.append("")
        
        # Add warnings if any
        if doc.warnings:
            lines.append("---")
            lines.append("*Parser warnings:*")
            for w in doc.warnings:
                lines.append(f"- {w}")
        
        return "\n".join(lines)


class JSONFormatter:
    """Converts ParsedDocument to structured JSON for programmatic use."""

    def format(self, doc: ParsedDocument) -> str:
        output = {
            "metadata": asdict(doc.metadata),
            "sections": [],
            "tables": [],
            "warnings": doc.warnings,
        }
        
        for section in doc.sections:
            output["sections"].append({
                "level": section.level,
                "title": section.title,
                "content": section.content,
                "word_count": len(section.content.split()) if section.content else 0,
            })
        
        for table in doc.tables:
            output["tables"].append({
                "caption": table.caption,
                "headers": table.headers,
                "rows": table.rows,
                "row_count": len(table.rows),
            })
        
        return json.dumps(output, indent=2, ensure_ascii=False)


class ChunkedFormatter:
    """Splits document into sized chunks for RAG pipelines."""

    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def format(self, doc: ParsedDocument) -> str:
        chunks = self._chunk_document(doc)
        output = {
            "metadata": asdict(doc.metadata),
            "chunk_config": {
                "chunk_size": self.chunk_size,
                "overlap": self.overlap,
                "total_chunks": len(chunks),
            },
            "chunks": chunks,
        }
        return json.dumps(output, indent=2, ensure_ascii=False)

    def _chunk_document(self, doc: ParsedDocument) -> list:
        chunks = []
        chunk_id = 0
        
        for section in doc.sections:
            content = section.content
            if not content:
                continue
            
            words = content.split()
            if len(words) <= self.chunk_size:
                chunks.append({
                    "id": chunk_id,
                    "section_title": section.title,
                    "section_level": section.level,
                    "content": content,
                    "word_count": len(words),
                })
                chunk_id += 1
            else:
                # Split into overlapping chunks
                start = 0
                while start < len(words):
                    end = min(start + self.chunk_size, len(words))
                    chunk_words = words[start:end]
                    chunks.append({
                        "id": chunk_id,
                        "section_title": section.title,
                        "section_level": section.level,
                        "content": " ".join(chunk_words),
                        "word_count": len(chunk_words),
                    })
                    chunk_id += 1
                    start = end - self.overlap
                    if start >= len(words):
                        break
        
        return chunks


# --- Main Application ---

SUPPORTED_EXTENSIONS = {
    ".pdf": PDFParser,
    ".docx": DOCXParser,
    ".doc": DOCXParser,  # May work for some .doc files
    ".txt": TextParser,
    ".md": TextParser,
}


def get_parser(filepath: str):
    """Return appropriate parser for file type."""
    ext = Path(filepath).suffix.lower()
    parser_class = SUPPORTED_EXTENSIONS.get(ext)
    if parser_class is None:
        raise ValueError(f"Unsupported file type: {ext}\nSupported: {', '.join(SUPPORTED_EXTENSIONS.keys())}")
    return parser_class()


def process_file(filepath: str, output_format: str = "markdown",
                 chunk_size: int = 1000, overlap: int = 100,
                 include_metadata: bool = True) -> tuple[str, str]:
    """Process a single file and return (output_content, suggested_extension)."""
    parser = get_parser(filepath)
    doc = parser.parse(filepath)
    
    if output_format == "markdown":
        formatter = MarkdownFormatter()
        content = formatter.format(doc, include_metadata=include_metadata)
        return content, ".md"
    elif output_format == "json":
        formatter = JSONFormatter()
        content = formatter.format(doc)
        return content, ".json"
    elif output_format == "chunked":
        formatter = ChunkedFormatter(chunk_size=chunk_size, overlap=overlap)
        content = formatter.format(doc)
        return content, ".chunks.json"
    else:
        raise ValueError(f"Unknown format: {output_format}")


def process_directory(input_dir: str, output_dir: str, output_format: str = "markdown",
                      chunk_size: int = 1000, overlap: int = 100,
                      include_metadata: bool = True, recursive: bool = True) -> dict:
    """Process all supported files in a directory."""
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeRemainingColumn
    
    console = Console()
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Collect files
    files = []
    pattern = "**/*" if recursive else "*"
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(input_path.glob(f"{pattern}{ext}"))
        files.extend(input_path.glob(f"{pattern}{ext.upper()}"))
    
    # Deduplicate
    files = sorted(set(files))
    
    if not files:
        console.print(f"[yellow]No supported files found in {input_dir}[/yellow]")
        return {"processed": 0, "failed": 0, "files": []}
    
    console.print(f"\n[bold]Found {len(files)} files to process[/bold]\n")
    
    results = {"processed": 0, "failed": 0, "files": [], "total_words": 0, "total_chars_saved": 0}
    
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TimeRemainingColumn(),
        console=console,
    ) as progress:
        task = progress.add_task("Processing files...", total=len(files))
        
        for filepath in files:
            rel_path = filepath.relative_to(input_path)
            progress.update(task, description=f"Processing {rel_path.name}...")
            
            try:
                content, ext = process_file(
                    str(filepath), output_format, chunk_size, overlap, include_metadata
                )
                
                # Preserve directory structure
                out_file = output_path / rel_path.with_suffix(ext)
                out_file.parent.mkdir(parents=True, exist_ok=True)
                out_file.write_text(content, encoding="utf-8")
                
                original_size = filepath.stat().st_size
                new_size = len(content.encode("utf-8"))
                
                results["processed"] += 1
                results["files"].append({
                    "source": str(rel_path),
                    "output": str(out_file.relative_to(output_path)),
                    "original_bytes": original_size,
                    "output_bytes": new_size,
                    "reduction_pct": round((1 - new_size / max(original_size, 1)) * 100, 1),
                })
                results["total_chars_saved"] += max(0, original_size - new_size)
                
            except Exception as e:
                results["failed"] += 1
                results["files"].append({
                    "source": str(rel_path),
                    "error": str(e),
                })
                console.print(f"  [red]✗ {rel_path.name}: {e}[/red]")
            
            progress.advance(task)
    
    return results


def generate_manifest(results: dict, output_dir: str):
    """Generate a manifest file summarizing all processed documents."""
    manifest = {
        "generated_at": datetime.now().isoformat(),
        "summary": {
            "total_processed": results["processed"],
            "total_failed": results["failed"],
        },
        "files": results["files"],
    }
    
    manifest_path = Path(output_dir) / "_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return str(manifest_path)


def print_summary(results: dict, output_dir: str):
    """Print a formatted summary of processing results."""
    from rich.console import Console
    from rich.table import Table
    
    console = Console()
    
    console.print(f"\n[bold green]Processing Complete[/bold green]")
    console.print(f"  Processed: {results['processed']}  |  Failed: {results['failed']}\n")
    
    if results["files"]:
        table = Table(title="File Results")
        table.add_column("Source", style="cyan")
        table.add_column("Status", style="green")
        table.add_column("Original", justify="right")
        table.add_column("Output", justify="right")
        table.add_column("Reduction", justify="right")
        
        for f in results["files"]:
            if "error" in f:
                table.add_row(f["source"], "[red]FAILED[/red]", "", "", f["error"][:40])
            else:
                orig = f"{f['original_bytes']:,} B"
                out = f"{f['output_bytes']:,} B"
                reduction = f"{f['reduction_pct']}%"
                table.add_row(f["source"], "[green]OK[/green]", orig, out, reduction)
        
        console.print(table)
    
    console.print(f"\n  Output directory: [bold]{output_dir}[/bold]\n")


def main():
    parser = argparse.ArgumentParser(
        description="DocParser - Convert documents to structured data for LLM consumption",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert a single PDF to Markdown
  python docparser.py document.pdf

  # Convert all files in a folder to JSON
  python docparser.py ./documents/ -f json -o ./parsed/

  # Generate RAG-ready chunks
  python docparser.py ./documents/ -f chunked --chunk-size 500 --overlap 50

  # Process a single DOCX without metadata header
  python docparser.py report.docx --no-metadata
        """
    )
    
    parser.add_argument("input", help="Input file or directory path")
    parser.add_argument("-o", "--output", help="Output file or directory (default: ./parsed_output/)")
    parser.add_argument("-f", "--format", choices=["markdown", "json", "chunked"],
                        default="markdown", help="Output format (default: markdown)")
    parser.add_argument("--chunk-size", type=int, default=1000,
                        help="Words per chunk for chunked format (default: 1000)")
    parser.add_argument("--overlap", type=int, default=100,
                        help="Word overlap between chunks (default: 100)")
    parser.add_argument("--no-metadata", action="store_true",
                        help="Omit YAML frontmatter metadata in markdown output")
    parser.add_argument("--no-recursive", action="store_true",
                        help="Don't process subdirectories")
    parser.add_argument("--manifest", action="store_true",
                        help="Generate a JSON manifest of all processed files")
    
    args = parser.parse_args()
    input_path = Path(args.input)
    
    if not input_path.exists():
        print(f"Error: {args.input} does not exist", file=sys.stderr)
        sys.exit(1)
    
    from rich.console import Console
    console = Console()
    
    if input_path.is_file():
        # Single file mode
        output_path = args.output
        if output_path is None:
            content, ext = process_file(
                str(input_path), args.format, args.chunk_size, args.overlap,
                not args.no_metadata
            )
            output_path = str(input_path.with_suffix(ext))
        
        try:
            content, ext = process_file(
                str(input_path), args.format, args.chunk_size, args.overlap,
                not args.no_metadata
            )
            
            if output_path is None:
                output_path = str(input_path.with_suffix(ext))
            
            Path(output_path).write_text(content, encoding="utf-8")
            console.print(f"\n[green]✓[/green] Converted: {input_path.name} → {output_path}")
            
            orig_size = input_path.stat().st_size
            new_size = len(content.encode("utf-8"))
            console.print(f"  {orig_size:,} bytes → {new_size:,} bytes "
                         f"({(1 - new_size/max(orig_size,1))*100:.1f}% reduction)\n")
        except Exception as e:
            console.print(f"\n[red]Error:[/red] {e}")
            sys.exit(1)
    
    elif input_path.is_dir():
        # Directory mode
        output_dir = args.output or "./parsed_output"
        
        results = process_directory(
            str(input_path), output_dir, args.format,
            args.chunk_size, args.overlap,
            not args.no_metadata, not args.no_recursive
        )
        
        print_summary(results, output_dir)
        
        if args.manifest:
            manifest_path = generate_manifest(results, output_dir)
            console.print(f"  Manifest: [bold]{manifest_path}[/bold]\n")
    else:
        print(f"Error: {args.input} is not a file or directory", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

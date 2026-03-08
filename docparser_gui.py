#!/usr/bin/env python3
"""
DocParser GUI - Document to Structured Data Converter
Cross-platform desktop application for converting PDFs and Office documents
into clean, structured Markdown/JSON optimized for LLM consumption.

Works on: Windows, macOS, Linux
Dependencies: pip install customtkinter pymupdf python-docx pyyaml
"""

import customtkinter as ctk
import hashlib
import json
import os
import re
import sys
import threading
import time
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox
from typing import Optional

# ============================================================================
# Document Models
# ============================================================================

@dataclass
class TableData:
    headers: list[str]
    rows: list[list[str]]
    caption: str = ""

    def to_markdown(self) -> str:
        if not self.headers and not self.rows:
            return ""
        lines = []
        if self.caption:
            lines.append(f"*{self.caption}*\n")
        if self.headers:
            hdrs = self.headers
        elif self.rows:
            hdrs = [f"Col {i+1}" for i in range(len(self.rows[0]))]
        else:
            return ""
        lines.append("| " + " | ".join(hdrs) + " |")
        lines.append("| " + " | ".join(["---"] * len(hdrs)) + " |")
        for row in self.rows:
            padded = row + [""] * (len(hdrs) - len(row))
            lines.append("| " + " | ".join(padded[:len(hdrs)]) + " |")
        return "\n".join(lines)


@dataclass
class Section:
    level: int
    title: str
    content: str
    subsections: list = field(default_factory=list)
    tables: list = field(default_factory=list)


@dataclass
class DocumentMetadata:
    filename: str
    filepath: str
    file_type: str
    file_size_bytes: int
    page_count: int = 0
    title: str = ""
    author: str = ""
    created_date: str = ""
    modified_date: str = ""
    subject: str = ""
    keywords: list = field(default_factory=list)
    word_count: int = 0
    char_count: int = 0
    sha256: str = ""
    parsed_at: str = ""


@dataclass
class ParsedDocument:
    metadata: DocumentMetadata
    sections: list
    raw_text: str = ""
    tables: list = field(default_factory=list)
    warnings: list = field(default_factory=list)


# ============================================================================
# Parsers
# ============================================================================

class PDFParser:
    def parse(self, filepath: str) -> ParsedDocument:
        import fitz
        doc = fitz.open(filepath)
        path = Path(filepath)
        meta = doc.metadata or {}
        file_stat = path.stat()
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()

        all_text_blocks = []
        all_tables = []
        warnings = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("dict", flags=11)["blocks"]  # TEXT_PRESERVE_WHITESPACE
            for block in blocks:
                if block["type"] == 0:
                    for line in block.get("lines", []):
                        text = ""
                        max_font_size = 0
                        is_bold = False
                        for span in line.get("spans", []):
                            text += span["text"]
                            max_font_size = max(max_font_size, span["size"])
                            if "bold" in span.get("font", "").lower() or "Bold" in span.get("font", ""):
                                is_bold = True
                        text = text.strip()
                        if text:
                            all_text_blocks.append({
                                "text": text, "font_size": max_font_size,
                                "is_bold": is_bold, "page": page_num + 1,
                                "y_pos": line["bbox"][1],
                            })
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table in tables.tables:
                        extracted = table.extract()
                        if extracted and len(extracted) > 1:
                            headers = [str(c).strip() if c else "" for c in extracted[0]]
                            rows = [[str(c).strip() if c else "" for c in r] for r in extracted[1:]]
                            all_tables.append(TableData(headers=headers, rows=rows,
                                                        caption=f"Table from page {page_num + 1}"))
            except Exception as e:
                warnings.append(f"Table extraction failed on page {page_num + 1}: {e}")

        sections = self._build_sections(all_text_blocks)
        raw_text = "\n".join(b["text"] for b in all_text_blocks)
        metadata = DocumentMetadata(
            filename=path.name, filepath=str(path.absolute()), file_type="pdf",
            file_size_bytes=file_stat.st_size, page_count=len(doc),
            title=meta.get("title", "") or path.stem, author=meta.get("author", ""),
            created_date=meta.get("creationDate", ""), modified_date=meta.get("modDate", ""),
            subject=meta.get("subject", ""),
            keywords=[k.strip() for k in meta.get("keywords", "").split(",") if k.strip()],
            word_count=len(raw_text.split()), char_count=len(raw_text),
            sha256=sha256, parsed_at=datetime.now().isoformat(),
        )
        doc.close()
        return ParsedDocument(metadata=metadata, sections=sections, raw_text=raw_text,
                              tables=all_tables, warnings=warnings)

    def _build_sections(self, blocks: list) -> list:
        if not blocks:
            return []
        font_sizes = [b["font_size"] for b in blocks]
        if not font_sizes:
            return []
        avg_size = sum(font_sizes) / len(font_sizes)
        max_size = max(font_sizes)
        h1_threshold = avg_size * 1.5
        h2_threshold = avg_size * 1.25
        h3_threshold = avg_size * 1.1

        sections = []
        current_section = None
        current_content_lines = []

        for block in blocks:
            text, size, is_bold = block["text"], block["font_size"], block["is_bold"]
            heading_level = 0
            if size >= h1_threshold or (size == max_size and is_bold):
                heading_level = 1
            elif size >= h2_threshold and is_bold:
                heading_level = 2
            elif (size >= h3_threshold and is_bold) or (is_bold and size > avg_size):
                heading_level = 3

            if heading_level > 0 and len(text) < 200:
                if current_section or current_content_lines:
                    if current_section is None:
                        current_section = Section(level=1, title="", content="")
                    current_section.content = "\n".join(current_content_lines).strip()
                    sections.append(current_section)
                    current_content_lines = []
                current_section = Section(level=heading_level, title=text, content="")
            else:
                current_content_lines.append(text)

        if current_content_lines:
            if current_section is None:
                current_section = Section(level=1, title="Document Content", content="")
            current_section.content = "\n".join(current_content_lines).strip()
            sections.append(current_section)
        elif current_section:
            sections.append(current_section)
        return sections


class DOCXParser:
    def parse(self, filepath: str) -> ParsedDocument:
        from docx import Document
        doc = Document(filepath)
        path = Path(filepath)
        file_stat = path.stat()
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
        props = doc.core_properties

        sections = []
        tables = []
        raw_lines = []
        current_section = None
        current_content_lines = []
        warnings = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                raw_lines.append(text)
                style_name = (para.style.name or "").lower() if para.style else ""

                heading_level = 0
                if "heading" in style_name:
                    try:
                        heading_level = int(re.search(r'\d+', style_name).group())
                    except (AttributeError, ValueError):
                        heading_level = 1
                elif style_name == "title":
                    heading_level = 1
                elif style_name == "subtitle":
                    heading_level = 2

                if heading_level > 0:
                    if current_section or current_content_lines:
                        if current_section is None:
                            current_section = Section(level=1, title="", content="")
                        current_section.content = "\n".join(current_content_lines).strip()
                        sections.append(current_section)
                        current_content_lines = []
                    current_section = Section(level=heading_level, title=text, content="")
                else:
                    if style_name.startswith("list"):
                        text = f"- {text}"
                    current_content_lines.append(text)

            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        try:
                            headers = []
                            rows = []
                            for i, row in enumerate(table.rows):
                                cells = [cell.text.strip() for cell in row.cells]
                                if i == 0:
                                    headers = cells
                                else:
                                    rows.append(cells)
                            td = TableData(headers=headers, rows=rows)
                            tables.append(td)
                            current_content_lines.append("")
                            current_content_lines.append(td.to_markdown())
                            current_content_lines.append("")
                        except Exception as e:
                            warnings.append(f"Table extraction error: {e}")
                        break

        if current_content_lines:
            if current_section is None:
                current_section = Section(level=1, title="Document Content", content="")
            current_section.content = "\n".join(current_content_lines).strip()
            sections.append(current_section)
        elif current_section:
            sections.append(current_section)

        raw_text = "\n".join(raw_lines)
        metadata = DocumentMetadata(
            filename=path.name, filepath=str(path.absolute()), file_type="docx",
            file_size_bytes=file_stat.st_size, page_count=0,
            title=props.title or path.stem, author=props.author or "",
            created_date=str(props.created) if props.created else "",
            modified_date=str(props.modified) if props.modified else "",
            subject=props.subject or "",
            keywords=[k.strip() for k in (props.keywords or "").split(",") if k.strip()],
            word_count=len(raw_text.split()), char_count=len(raw_text),
            sha256=sha256, parsed_at=datetime.now().isoformat(),
        )
        return ParsedDocument(metadata=metadata, sections=sections, raw_text=raw_text,
                              tables=tables, warnings=warnings)


class TextParser:
    def parse(self, filepath: str) -> ParsedDocument:
        path = Path(filepath)
        file_stat = path.stat()
        content = path.read_text(encoding="utf-8", errors="replace")
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
        ext = path.suffix.lower()
        sections = []

        if ext == ".md":
            current_section = None
            current_lines = []
            for line in content.split("\n"):
                heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
                if heading_match:
                    if current_section or current_lines:
                        if current_section is None:
                            current_section = Section(level=1, title="", content="")
                        current_section.content = "\n".join(current_lines).strip()
                        sections.append(current_section)
                        current_lines = []
                    level = len(heading_match.group(1))
                    current_section = Section(level=level, title=heading_match.group(2).strip(), content="")
                else:
                    current_lines.append(line)
            if current_lines:
                if current_section is None:
                    current_section = Section(level=1, title="Document Content", content="")
                current_section.content = "\n".join(current_lines).strip()
                sections.append(current_section)
        else:
            sections.append(Section(level=1, title=path.stem, content=content.strip()))

        metadata = DocumentMetadata(
            filename=path.name, filepath=str(path.absolute()), file_type=ext.lstrip("."),
            file_size_bytes=file_stat.st_size, title=path.stem,
            word_count=len(content.split()), char_count=len(content),
            sha256=sha256, parsed_at=datetime.now().isoformat(),
        )
        return ParsedDocument(metadata=metadata, sections=sections, raw_text=content)


# ============================================================================
# Formatters
# ============================================================================

class MarkdownFormatter:
    def format(self, doc: ParsedDocument, include_metadata: bool = True) -> str:
        lines = []
        if include_metadata:
            lines.extend([
                "---", f'title: "{doc.metadata.title}"',
                f'source_file: "{doc.metadata.filename}"', f'file_type: {doc.metadata.file_type}',
            ])
            if doc.metadata.author:
                lines.append(f'author: "{doc.metadata.author}"')
            if doc.metadata.page_count:
                lines.append(f'pages: {doc.metadata.page_count}')
            lines.extend([
                f'word_count: {doc.metadata.word_count}',
                f'parsed_at: {doc.metadata.parsed_at}',
                f'sha256: {doc.metadata.sha256[:16]}...', "---", "",
            ])
        for section in doc.sections:
            if section.title:
                lines.extend([f"{'#' * section.level} {section.title}", ""])
            if section.content:
                lines.extend([re.sub(r'\n{3,}', '\n\n', section.content), ""])
        if doc.tables:
            has_tables_in_sections = any("| " in s.content for s in doc.sections if s.content)
            if not has_tables_in_sections:
                lines.extend(["## Extracted Tables", ""])
                for table in doc.tables:
                    if table.caption:
                        lines.append(f"### {table.caption}")
                    lines.extend([table.to_markdown(), ""])
        if doc.warnings:
            lines.append("---\n*Parser warnings:*")
            for w in doc.warnings:
                lines.append(f"- {w}")
        return "\n".join(lines)


class JSONFormatter:
    def format(self, doc: ParsedDocument) -> str:
        output = {
            "metadata": asdict(doc.metadata),
            "sections": [{"level": s.level, "title": s.title, "content": s.content,
                          "word_count": len(s.content.split()) if s.content else 0}
                         for s in doc.sections],
            "tables": [{"caption": t.caption, "headers": t.headers, "rows": t.rows,
                        "row_count": len(t.rows)} for t in doc.tables],
            "warnings": doc.warnings,
        }
        return json.dumps(output, indent=2, ensure_ascii=False)


class ChunkedFormatter:
    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def format(self, doc: ParsedDocument) -> str:
        chunks = []
        chunk_id = 0
        for section in doc.sections:
            if not section.content:
                continue
            words = section.content.split()
            if len(words) <= self.chunk_size:
                chunks.append({"id": chunk_id, "section_title": section.title,
                               "section_level": section.level, "content": section.content,
                               "word_count": len(words)})
                chunk_id += 1
            else:
                start = 0
                while start < len(words):
                    end = min(start + self.chunk_size, len(words))
                    chunk_words = words[start:end]
                    chunks.append({"id": chunk_id, "section_title": section.title,
                                   "section_level": section.level,
                                   "content": " ".join(chunk_words),
                                   "word_count": len(chunk_words)})
                    chunk_id += 1
                    start = end - self.overlap
                    if start >= len(words):
                        break
        output = {
            "metadata": asdict(doc.metadata),
            "chunk_config": {"chunk_size": self.chunk_size, "overlap": self.overlap,
                             "total_chunks": len(chunks)},
            "chunks": chunks,
        }
        return json.dumps(output, indent=2, ensure_ascii=False)


# ============================================================================
# Processing Engine
# ============================================================================

SUPPORTED_EXTENSIONS = {".pdf": PDFParser, ".docx": DOCXParser, ".doc": DOCXParser,
                        ".txt": TextParser, ".md": TextParser}

FORMAT_EXTENSIONS = {"Markdown (.md)": ".md", "JSON (.json)": ".json",
                     "Chunked JSON (.chunks.json)": ".chunks.json"}

FORMAT_KEYS = {"Markdown (.md)": "markdown", "JSON (.json)": "json",
               "Chunked JSON (.chunks.json)": "chunked"}


def process_single_file(filepath: str, output_format: str, output_dir: str,
                        chunk_size: int = 1000, overlap: int = 100,
                        include_metadata: bool = True) -> dict:
    """Process one file and return result dict."""
    path = Path(filepath)
    ext = path.suffix.lower()
    parser_class = SUPPORTED_EXTENSIONS.get(ext)
    if parser_class is None:
        return {"source": path.name, "error": f"Unsupported file type: {ext}"}

    try:
        parser = parser_class()
        doc = parser.parse(filepath)

        if output_format == "markdown":
            content = MarkdownFormatter().format(doc, include_metadata)
            out_ext = ".md"
        elif output_format == "json":
            content = JSONFormatter().format(doc)
            out_ext = ".json"
        elif output_format == "chunked":
            content = ChunkedFormatter(chunk_size, overlap).format(doc)
            out_ext = ".chunks.json"
        else:
            return {"source": path.name, "error": f"Unknown format: {output_format}"}

        out_path = Path(output_dir) / (path.stem + out_ext)
        # Handle duplicate filenames
        counter = 1
        while out_path.exists():
            out_path = Path(output_dir) / f"{path.stem}_{counter}{out_ext}"
            counter += 1

        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(content, encoding="utf-8")

        original_size = path.stat().st_size
        new_size = len(content.encode("utf-8"))
        reduction = round((1 - new_size / max(original_size, 1)) * 100, 1)

        return {
            "source": path.name,
            "output": out_path.name,
            "output_path": str(out_path),
            "original_bytes": original_size,
            "output_bytes": new_size,
            "reduction_pct": reduction,
            "word_count": doc.metadata.word_count,
            "warnings": doc.warnings,
        }
    except Exception as e:
        return {"source": path.name, "error": str(e)}


# ============================================================================
# GUI Application
# ============================================================================

class FileListItem(ctk.CTkFrame):
    """Individual file entry in the file list."""

    def __init__(self, parent, filepath: str, on_remove=None, **kwargs):
        super().__init__(parent, **kwargs)
        self.filepath = filepath
        self.configure(fg_color="transparent")

        path = Path(filepath)
        ext = path.suffix.lower()
        size = path.stat().st_size
        size_str = self._format_size(size)

        # File type indicator
        type_colors = {".pdf": "#E74C3C", ".docx": "#3498DB", ".doc": "#3498DB",
                       ".txt": "#95A5A6", ".md": "#2ECC71"}
        color = type_colors.get(ext, "#95A5A6")

        self.type_label = ctk.CTkLabel(self, text=ext.upper().lstrip("."), width=50,
                                       font=ctk.CTkFont(size=11, weight="bold"),
                                       text_color=color)
        self.type_label.pack(side="left", padx=(8, 4))

        # File info
        info_frame = ctk.CTkFrame(self, fg_color="transparent")
        info_frame.pack(side="left", fill="x", expand=True, padx=4)

        self.name_label = ctk.CTkLabel(info_frame, text=path.name, anchor="w",
                                       font=ctk.CTkFont(size=13))
        self.name_label.pack(anchor="w")

        self.detail_label = ctk.CTkLabel(info_frame, text=f"{size_str}  •  {path.parent}",
                                         anchor="w", font=ctk.CTkFont(size=11),
                                         text_color="gray")
        self.detail_label.pack(anchor="w")

        # Status indicator (shown after processing)
        self.status_label = ctk.CTkLabel(self, text="", width=80,
                                         font=ctk.CTkFont(size=11))
        self.status_label.pack(side="right", padx=4)

        # Remove button
        self.remove_btn = ctk.CTkButton(self, text="✕", width=30, height=28,
                                        font=ctk.CTkFont(size=14),
                                        fg_color="transparent",
                                        hover_color=("gray85", "gray25"),
                                        text_color=("gray40", "gray60"),
                                        command=lambda: on_remove(self) if on_remove else None)
        self.remove_btn.pack(side="right", padx=(0, 4))

    def set_status(self, text: str, color: str = "gray"):
        self.status_label.configure(text=text, text_color=color)

    def set_processing(self):
        self.status_label.configure(text="Processing...", text_color="orange")
        self.remove_btn.configure(state="disabled")

    def disable_remove(self):
        self.remove_btn.configure(state="disabled")

    def enable_remove(self):
        self.remove_btn.configure(state="normal")

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"


class DocParserApp(ctk.CTk):
    """Main application window."""

    APP_NAME = "DocParser"
    APP_VERSION = "1.0.0"
    WINDOW_WIDTH = 820
    WINDOW_HEIGHT = 720

    def __init__(self):
        super().__init__()

        self.title(f"{self.APP_NAME} — Document to Structured Data")
        self.geometry(f"{self.WINDOW_WIDTH}x{self.WINDOW_HEIGHT}")
        self.minsize(700, 600)

        # State
        self.file_items: list[FileListItem] = []
        self.is_processing = False
        self.output_dir = str(Path.home() / "DocParser_Output")

        # Theme
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        self._build_ui()

    def _build_ui(self):
        # ── Header ──
        header = ctk.CTkFrame(self, height=60, fg_color=("gray92", "gray14"))
        header.pack(fill="x", padx=0, pady=0)
        header.pack_propagate(False)

        ctk.CTkLabel(header, text="⚡ DocParser",
                     font=ctk.CTkFont(size=22, weight="bold")).pack(side="left", padx=16)
        ctk.CTkLabel(header, text="Convert documents to LLM-ready structured data",
                     font=ctk.CTkFont(size=13),
                     text_color=("gray40", "gray60")).pack(side="left", padx=8)

        theme_btn = ctk.CTkButton(header, text="🌓", width=32, height=32,
                                  fg_color="transparent", hover_color=("gray80", "gray25"),
                                  command=self._toggle_theme, font=ctk.CTkFont(size=16))
        theme_btn.pack(side="right", padx=12)

        # ── Main content area ──
        main = ctk.CTkFrame(self, fg_color="transparent")
        main.pack(fill="both", expand=True, padx=16, pady=(12, 8))

        # ── File list section ──
        file_section = ctk.CTkFrame(main, fg_color="transparent")
        file_section.pack(fill="both", expand=True)

        # File list header with buttons
        file_header = ctk.CTkFrame(file_section, fg_color="transparent")
        file_header.pack(fill="x", pady=(0, 6))

        ctk.CTkLabel(file_header, text="Files to Convert",
                     font=ctk.CTkFont(size=15, weight="bold")).pack(side="left")

        self.file_count_label = ctk.CTkLabel(file_header, text="0 files",
                                             font=ctk.CTkFont(size=12),
                                             text_color=("gray40", "gray60"))
        self.file_count_label.pack(side="left", padx=12)

        clear_btn = ctk.CTkButton(file_header, text="Clear All", width=80, height=28,
                                  fg_color="transparent", border_width=1,
                                  border_color=("gray70", "gray40"),
                                  hover_color=("gray85", "gray25"),
                                  text_color=("gray30", "gray70"),
                                  font=ctk.CTkFont(size=12),
                                  command=self._clear_files)
        clear_btn.pack(side="right", padx=4)

        add_folder_btn = ctk.CTkButton(file_header, text="+ Add Folder", width=100, height=28,
                                       fg_color="transparent", border_width=1,
                                       border_color=("gray70", "gray40"),
                                       hover_color=("gray85", "gray25"),
                                       text_color=("gray30", "gray70"),
                                       font=ctk.CTkFont(size=12),
                                       command=self._add_folder)
        add_folder_btn.pack(side="right", padx=4)

        add_btn = ctk.CTkButton(file_header, text="+ Add Files", width=100, height=28,
                                fg_color=("gray75", "gray30"),
                                hover_color=("gray65", "gray35"),
                                font=ctk.CTkFont(size=12),
                                command=self._add_files)
        add_btn.pack(side="right", padx=4)

        # Scrollable file list
        self.file_list_frame = ctk.CTkScrollableFrame(file_section, fg_color=("gray95", "gray17"),
                                                       corner_radius=8)
        self.file_list_frame.pack(fill="both", expand=True)

        # Drop zone placeholder
        self.drop_placeholder = ctk.CTkLabel(
            self.file_list_frame,
            text="Click 'Add Files' or 'Add Folder' to select documents\n\n"
                 "Supported: PDF, DOCX, TXT, MD",
            font=ctk.CTkFont(size=13),
            text_color=("gray50", "gray50"),
            justify="center",
        )
        self.drop_placeholder.pack(expand=True, pady=40)

        # ── Settings section ──
        settings_frame = ctk.CTkFrame(main, fg_color=("gray92", "gray17"), corner_radius=8)
        settings_frame.pack(fill="x", pady=(12, 0))

        settings_inner = ctk.CTkFrame(settings_frame, fg_color="transparent")
        settings_inner.pack(fill="x", padx=16, pady=12)

        # Row 1: Format and Options
        row1 = ctk.CTkFrame(settings_inner, fg_color="transparent")
        row1.pack(fill="x", pady=(0, 8))

        # Format selector
        fmt_frame = ctk.CTkFrame(row1, fg_color="transparent")
        fmt_frame.pack(side="left")

        ctk.CTkLabel(fmt_frame, text="Output Format",
                     font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w")
        self.format_var = ctk.StringVar(value="Markdown (.md)")
        self.format_menu = ctk.CTkOptionMenu(
            fmt_frame, variable=self.format_var, width=220, height=32,
            values=list(FORMAT_KEYS.keys()),
            command=self._on_format_change,
            font=ctk.CTkFont(size=12),
        )
        self.format_menu.pack(anchor="w", pady=(4, 0))

        # Chunk settings (shown only for chunked format)
        self.chunk_frame = ctk.CTkFrame(row1, fg_color="transparent")

        ctk.CTkLabel(self.chunk_frame, text="Chunk Size (words)",
                     font=ctk.CTkFont(size=12)).pack(side="left", padx=(20, 4))
        self.chunk_size_var = ctk.StringVar(value="1000")
        ctk.CTkEntry(self.chunk_frame, textvariable=self.chunk_size_var,
                     width=70, height=32).pack(side="left")

        ctk.CTkLabel(self.chunk_frame, text="Overlap",
                     font=ctk.CTkFont(size=12)).pack(side="left", padx=(12, 4))
        self.overlap_var = ctk.StringVar(value="100")
        ctk.CTkEntry(self.chunk_frame, textvariable=self.overlap_var,
                     width=60, height=32).pack(side="left")

        # Metadata toggle
        self.metadata_var = ctk.BooleanVar(value=True)
        self.metadata_check = ctk.CTkCheckBox(row1, text="Include metadata header",
                                              variable=self.metadata_var,
                                              font=ctk.CTkFont(size=12))
        self.metadata_check.pack(side="right", padx=8)

        # Row 2: Output directory
        row2 = ctk.CTkFrame(settings_inner, fg_color="transparent")
        row2.pack(fill="x")

        ctk.CTkLabel(row2, text="Output Location",
                     font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w")

        dir_row = ctk.CTkFrame(row2, fg_color="transparent")
        dir_row.pack(fill="x", pady=(4, 0))

        self.output_dir_var = ctk.StringVar(value=self.output_dir)
        self.dir_entry = ctk.CTkEntry(dir_row, textvariable=self.output_dir_var,
                                      height=32, font=ctk.CTkFont(size=12))
        self.dir_entry.pack(side="left", fill="x", expand=True)

        browse_btn = ctk.CTkButton(dir_row, text="Browse", width=80, height=32,
                                   font=ctk.CTkFont(size=12),
                                   command=self._browse_output)
        browse_btn.pack(side="right", padx=(8, 0))

        # ── Bottom bar with process button ──
        bottom = ctk.CTkFrame(self, height=70, fg_color="transparent")
        bottom.pack(fill="x", padx=16, pady=(8, 12))
        bottom.pack_propagate(False)

        # Progress bar (hidden by default)
        self.progress_frame = ctk.CTkFrame(bottom, fg_color="transparent")

        self.progress_label = ctk.CTkLabel(self.progress_frame, text="Processing...",
                                           font=ctk.CTkFont(size=12))
        self.progress_label.pack(anchor="w")

        self.progress_bar = ctk.CTkProgressBar(self.progress_frame, height=8)
        self.progress_bar.pack(fill="x", pady=(4, 0))
        self.progress_bar.set(0)

        # Stats label (shown after processing)
        self.stats_label = ctk.CTkLabel(bottom, text="", font=ctk.CTkFont(size=12),
                                        text_color=("gray40", "gray60"))
        self.stats_label.pack(side="left", padx=4)

        # Process button
        self.process_btn = ctk.CTkButton(
            bottom, text="▶  Process Files", width=180, height=42,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color=("#2563EB", "#2563EB"),
            hover_color=("#1D4ED8", "#1D4ED8"),
            command=self._start_processing,
        )
        self.process_btn.pack(side="right")

        # Open output folder button (hidden by default)
        self.open_folder_btn = ctk.CTkButton(
            bottom, text="📂 Open Output Folder", width=160, height=36,
            font=ctk.CTkFont(size=12),
            fg_color="transparent", border_width=1,
            border_color=("gray70", "gray40"),
            hover_color=("gray85", "gray25"),
            text_color=("gray30", "gray70"),
            command=self._open_output_folder,
        )

    # ── Theme toggle ──
    def _toggle_theme(self):
        current = ctk.get_appearance_mode()
        ctk.set_appearance_mode("light" if current == "Dark" else "dark")

    # ── File management ──
    def _add_files(self):
        filetypes = [
            ("All Supported", "*.pdf *.docx *.doc *.txt *.md"),
            ("PDF Files", "*.pdf"),
            ("Word Documents", "*.docx *.doc"),
            ("Text Files", "*.txt *.md"),
        ]
        files = filedialog.askopenfilenames(title="Select Files", filetypes=filetypes)
        if files:
            self._add_file_paths(files)

    def _add_folder(self):
        folder = filedialog.askdirectory(title="Select Folder")
        if folder:
            files = []
            for ext in SUPPORTED_EXTENSIONS:
                files.extend(Path(folder).rglob(f"*{ext}"))
                files.extend(Path(folder).rglob(f"*{ext.upper()}"))
            self._add_file_paths([str(f) for f in sorted(set(files))])

    def _add_file_paths(self, paths):
        existing = {item.filepath for item in self.file_items}
        added = 0
        for fp in paths:
            if fp not in existing:
                self.drop_placeholder.pack_forget()
                item = FileListItem(self.file_list_frame, fp,
                                    on_remove=self._remove_file)
                item.pack(fill="x", pady=2, padx=4)
                self.file_items.append(item)
                added += 1
        self._update_file_count()

    def _remove_file(self, item: FileListItem):
        if self.is_processing:
            return
        item.destroy()
        self.file_items.remove(item)
        self._update_file_count()
        if not self.file_items:
            self.drop_placeholder.pack(expand=True, pady=40)

    def _clear_files(self):
        if self.is_processing:
            return
        for item in self.file_items:
            item.destroy()
        self.file_items.clear()
        self.drop_placeholder.pack(expand=True, pady=40)
        self._update_file_count()
        self.stats_label.configure(text="")
        self.open_folder_btn.pack_forget()

    def _update_file_count(self):
        n = len(self.file_items)
        self.file_count_label.configure(text=f"{n} file{'s' if n != 1 else ''}")

    # ── Settings ──
    def _on_format_change(self, value):
        if "Chunked" in value:
            self.chunk_frame.pack(side="left")
        else:
            self.chunk_frame.pack_forget()

    def _browse_output(self):
        folder = filedialog.askdirectory(title="Select Output Folder")
        if folder:
            self.output_dir_var.set(folder)

    def _open_output_folder(self):
        output_dir = self.output_dir_var.get()
        if os.path.isdir(output_dir):
            if sys.platform == "win32":
                os.startfile(output_dir)
            elif sys.platform == "darwin":
                os.system(f'open "{output_dir}"')
            else:
                os.system(f'xdg-open "{output_dir}" 2>/dev/null &')

    # ── Processing ──
    def _start_processing(self):
        if self.is_processing or not self.file_items:
            if not self.file_items:
                messagebox.showinfo("No Files", "Add files to process first.")
            return

        output_dir = self.output_dir_var.get().strip()
        if not output_dir:
            messagebox.showerror("Error", "Please set an output directory.")
            return

        format_key = FORMAT_KEYS[self.format_var.get()]

        try:
            chunk_size = int(self.chunk_size_var.get())
            overlap = int(self.overlap_var.get())
        except ValueError:
            messagebox.showerror("Error", "Chunk size and overlap must be numbers.")
            return

        include_metadata = self.metadata_var.get()

        # Lock UI
        self.is_processing = True
        self.process_btn.configure(state="disabled", text="Processing...")
        self.stats_label.configure(text="")
        self.open_folder_btn.pack_forget()

        # Show progress
        self.progress_frame.pack(side="left", fill="x", expand=True, padx=(0, 12))
        self.progress_bar.set(0)

        for item in self.file_items:
            item.set_status("")
            item.disable_remove()

        # Run in background thread
        thread = threading.Thread(
            target=self._process_worker,
            args=(format_key, output_dir, chunk_size, overlap, include_metadata),
            daemon=True,
        )
        thread.start()

    def _process_worker(self, format_key, output_dir, chunk_size, overlap, include_metadata):
        total = len(self.file_items)
        processed = 0
        failed = 0
        total_original = 0
        total_output = 0

        for i, item in enumerate(self.file_items):
            # Update UI from main thread
            self.after(0, item.set_processing)
            self.after(0, lambda v=f"Processing {i+1}/{total}...": self.progress_label.configure(text=v))

            result = process_single_file(
                item.filepath, format_key, output_dir, chunk_size, overlap, include_metadata
            )

            if "error" in result:
                failed += 1
                self.after(0, lambda it=item, r=result: it.set_status(f"✗ {r['error'][:30]}", "#E74C3C"))
            else:
                processed += 1
                total_original += result["original_bytes"]
                total_output += result["output_bytes"]
                reduction = result["reduction_pct"]
                self.after(0, lambda it=item, r=reduction: it.set_status(f"✓ {r}% smaller", "#2ECC71"))

            progress = (i + 1) / total
            self.after(0, lambda v=progress: self.progress_bar.set(v))

        # Done - update UI
        def finish():
            self.is_processing = False
            self.process_btn.configure(state="normal", text="▶  Process Files")
            self.progress_frame.pack_forget()

            for item in self.file_items:
                item.enable_remove()

            if total_original > 0:
                overall_reduction = round((1 - total_output / total_original) * 100, 1)
                stats = f"✓ {processed} processed"
                if failed:
                    stats += f"  •  ✗ {failed} failed"
                stats += f"  •  {overall_reduction}% overall reduction"
                self.stats_label.configure(text=stats)
            elif failed:
                self.stats_label.configure(text=f"✗ {failed} file(s) failed")

            self.open_folder_btn.pack(side="right", padx=(8, 0))

            if processed > 0 and failed == 0:
                self.after(100, lambda: messagebox.showinfo(
                    "Complete",
                    f"Successfully converted {processed} file(s).\n"
                    f"Output saved to:\n{output_dir}"
                ))

        self.after(0, finish)


# ============================================================================
# Entry Point
# ============================================================================

def main():
    app = DocParserApp()
    app.mainloop()


if __name__ == "__main__":
    main()
